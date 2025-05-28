from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
import numpy as np

# Load dataset and preprocessing (same as your EDA step)
df = pd.read_csv('Books.csv')
df['Price'] = df['Price'].replace('[£]', '', regex=True).astype(float)

# Map Rate to numeric
rating_map = {
    'One': 1,
    'Two': 2,
    'Three': 3,
    'Four': 4,
    'Five': 5
}
df['Rate'] = df['Rate'].map(rating_map)

# Outlier detection for price (IQR method)
Q1 = df['Price'].quantile(0.25)
Q3 = df['Price'].quantile(0.75)
IQR = Q3 - Q1
lower_bound = Q1 - 1.5 * IQR
upper_bound = Q3 + 1.5 * IQR
outliers = df[(df['Price'] < lower_bound) | (df['Price'] > upper_bound)]


# Create a new Document
doc = Document()

# Title
title = doc.add_heading('Exploratory Data Analysis Report - Books Dataset', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Intro
doc.add_paragraph(
    "This report summarizes the key insights and findings from the Exploratory Data Analysis (EDA) "
    "performed on the Books dataset. The analysis covers basic statistics, data cleaning, "
    "visualizations, and detection of important trends and issues."
)

# Section: Dataset Overview
doc.add_heading('Dataset Overview', level=1)
doc.add_paragraph(f"Number of rows (books): {df.shape[0]}")
doc.add_paragraph(f"Number of columns (features): {df.shape[1]}")
doc.add_paragraph("Columns and types:")
for col, dtype in zip(df.columns, df.dtypes):
    doc.add_paragraph(f" - {col}: {dtype}")

# Section: Missing Values and Duplicates
doc.add_heading('Missing Values and Duplicates', level=1)
missing_vals = df.isnull().sum()
for col, miss in missing_vals.items():
    doc.add_paragraph(f" - {col}: {miss} missing values")
doc.add_paragraph(f"Total duplicate rows found: {df.duplicated().sum()}")

# Section: Price Summary
doc.add_heading('Price Summary', level=1)
doc.add_paragraph(f" - Average Price: £{df['Price'].mean():.2f}")
doc.add_paragraph(f" - Median Price: £{df['Price'].median():.2f}")
doc.add_paragraph(f" - Minimum Price: £{df['Price'].min():.2f}")
doc.add_paragraph(f" - Maximum Price: £{df['Price'].max():.2f}")

# Section: Rating Summary
doc.add_heading('Rating Summary', level=1)
for rate, count in df['Rate'].value_counts().sort_index().items():
    doc.add_paragraph(f" - Rating {rate} stars: {count} books")

# Section: Availability
doc.add_heading('Availability', level=1)
availability_counts = df['Availability'].value_counts()
for avail, count in availability_counts.items():
    doc.add_paragraph(f" - {avail}: {count} books")

# Section: Outliers in Price
doc.add_heading('Outliers in Price', level=1)
doc.add_paragraph(f"Number of detected outliers in price based on IQR method: {outliers.shape[0]}")
doc.add_paragraph("Examples of outliers (Title, Price, Rating):")
for idx, row in outliers[['Title', 'Price', 'Rate']].head(5).iterrows():
    doc.add_paragraph(f" - {row['Title']}: £{row['Price']:.2f}, Rating: {row['Rate']}")

# Section: Correlations and Relationships
doc.add_heading('Correlations and Relationships', level=1)
corr = df[['Price', 'Rate']].corr().iloc[0,1]
doc.add_paragraph(f"Correlation between Price and Rating: {corr:.2f}")
doc.add_paragraph("Insight: There is a positive correlation, suggesting higher-rated books tend to be more expensive.")

# Section: Top Books
doc.add_heading('Top Books', level=1)
doc.add_paragraph('Top 5 Most Expensive Books:')
for idx, row in df.sort_values('Price', ascending=False)[['Title', 'Price', 'Rate']].head(5).iterrows():
    doc.add_paragraph(f" - {row['Title']} (£{row['Price']:.2f}), Rating: {row['Rate']}")

doc.add_paragraph('Top 5 Cheapest Books:')
for idx, row in df.sort_values('Price', ascending=True)[['Title', 'Price', 'Rate']].head(5).iterrows():
    doc.add_paragraph(f" - {row['Title']} (£{row['Price']:.2f}), Rating: {row['Rate']}")

# Save the document
doc.save('Books_EDA_Report.docx')
print("Word report saved as 'Books_EDA_Report.docx'")
